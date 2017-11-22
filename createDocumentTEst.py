from docx import Document
from docx.shared import Inches
from makeRandomString import *

document = Document()

document.add_heading('Bad baby',0)

p = document.add_paragraph(makeRandomTyson(2000))
p.add_run(makeRandomStringLower(5000)).bold = True
p.add_run(makeRandomStringLower(6000))
p.add_run(makeRandomTyson(9)).italic = True

document.add_heading('READ THIS FIRST, level 1', level=1)
document.add_paragraph('then read this', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

##document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
##for item in recordset:
##        row_cells = table.add_rows().cells
##        row_cells[0].text = str(item.qty)
##        row_cells[1].text = str(item.id)
##        row_cells[2].text = item.desc

document.add_page_break()
document.save('demo.docx')
